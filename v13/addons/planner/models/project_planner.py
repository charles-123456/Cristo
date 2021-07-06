# -*- coding: utf-8 -*-
from odoo import models, fields, api, _
from odoo.exceptions import UserError, AccessError, Warning
from lxml import etree
from docx import Document
from docx.shared import Pt, Inches
import base64
from io import BytesIO
import datetime
# The below package is used to restart the numbering system in document paragraph
# https://pypi.org/project/python-docx-whtsky/

class ProjectPlan(models.Model):
    _name = "project.plan"
    _description = 'Project Plan'
    
    @api.model
    def fields_view_get(self, view_id=None, view_type='form', toolbar=False, submenu=False):
        res = super(ProjectPlan, self).fields_view_get(view_id=view_id, view_type=view_type, toolbar=toolbar, submenu=submenu)
        doc = etree.XML(res['arch'])
        current_user = self.env.user
        section1_config_id = self.env['plan.section.config'].search([('section_type','=','section 1'),('user_ids','=',current_user.id)])
        if view_type == 'form':
            if section1_config_id:
                for node in doc.xpath("//field[@name='section1']"):
                    node.set('string', section1_config_id.name)
                for node in doc.xpath("//notebook/page[2]/div[1]/h3"):
                    node.set('text', section1_config_id.name)
                for node in doc.xpath("//button[@id='sec1']"):
                    node.set('string', section1_config_id.name)
            else:
                for node in doc.xpath("//notebook/page[2]/div[1]/h3"):
                    node.set('text', 'Dimension')
                    
            section2_config_id = self.env['plan.section.config'].search([('section_type','=','section 2'),('user_ids','=',current_user.id)])
            if section2_config_id:
                for node in doc.xpath("//field[@name='section2']"):
                    node.set('string', section2_config_id.name)
                for node in doc.xpath("//notebook/page[2]/div[2]/h3"):
                    node.set('text', section2_config_id.name)
                for node in doc.xpath("//button[@id='sec2']"):
                    node.set('string', section2_config_id.name)
            else:
                for node in doc.xpath("//notebook/page[2]/div[2]/h3"):
                    node.set('text', 'Need')
                        
            section3_config_id = self.env['plan.section.config'].search([('section_type','=','section 3'),('user_ids','=',current_user.id)])
            if section3_config_id:
                for node in doc.xpath("//field[@name='section3']"):
                    node.set('string', section3_config_id.name)
                for node in doc.xpath("//notebook/page[2]/div[3]/h3"):
                    node.set('text', section3_config_id.name)
                for node in doc.xpath("//button[@id='sec3']"):
                    node.set('string', section3_config_id.name)
            else:
                for node in doc.xpath("//notebook/page[2]/div[3]/h3"):
                    node.set('text', 'Objective')
                      
            section4_config_id = self.env['plan.section.config'].search([('section_type','=','section 4'),('user_ids','=',current_user.id)])
            if section4_config_id:
                for node in doc.xpath("//field[@name='section4']"):
                    node.set('string', section4_config_id.name)
                for node in doc.xpath("//notebook/page[2]/div[4]/h3"):
                    node.set('text', section4_config_id.name)
                for node in doc.xpath("//button[@id='sec4']"):
                    node.set('string', section4_config_id.name)
            else:
                for node in doc.xpath("//notebook/page[2]/div[4]/h3"):
                    node.set('text', 'Strategy')
        res['arch'] = etree.tostring(doc, encoding='unicode')
        return res
    
    @api.constrains('start_date','end_date')
    def validate_date(self):
        if self.start_date and self.end_date:
            if self.start_date > self.end_date:
                raise UserError(_("Start date must be less than the End date."))
    
    @api.constrains('section1','section2','section3','section4')
    def validate_date(self):
        if self.section1 == 'no' and self.section2 == 'no' and self.section3 == 'no' and self.section4 == 'no':
            raise UserError(_("Please choose atleast one section."))
    
    def _compute_plan_activity(self):
        activities_count = 0
        for rec in self:
            self.activities_count = self.env['plan.activity'].search_count([('plan_id','=', rec.id)])
        
    def _compute_expenditure_amt(self):
        self.actual_expenditure = 0
        for rec in self:
            expenditures = self.env['plan.expenditure'].search([('plan_id','=', rec.id)])
            for expenditure in expenditures:
                rec.actual_expenditure +=  expenditure.amount
    
    @api.constrains('star_date','end_date')
    def date_validation(self):
        if self.end_date < self.start_date:
            raise UserError(_("End date to should not be lesser than start Date."))
           
    name = fields.Char(string="Plan Name", copy=False)
    category_id = fields.Many2one('plan.category', string="Plan Category")
    plan_type = fields.Selection([('Strategic','Strategic'),('Operational','Operational')], string="Plan Type")
    description = fields.Html(string="Description", copy=False)
    start_date = fields.Date(string="Start Date", copy=False)
    end_date = fields.Date(string="End Date", copy=False)
    responsible_id = fields.Many2one('res.partner', string="Primarily Responsible")
    section1 = fields.Selection([('yes','Yes'),('no','No')], string="Section 1", default='yes')
    section2 = fields.Selection([('yes','Yes'),('no','No')], string="Section 2", default='no')
    section3 = fields.Selection([('yes','Yes'),('no','No')], string="Section 3", default='no')
    section4 = fields.Selection([('yes','Yes'),('no','No')], string="Section 4", default='no')
    activity_type_ids = fields.Many2many('plan.main.activity.type', string="Activity Types", copy=False)
    activities_count = fields.Integer(string="No. of Activities", compute='_compute_plan_activity')
    estimated_expenditure = fields.Float(string="Estimated Expenditure", copy=False)
    actual_expenditure = fields.Float(string="Actual Expenditure", compute='_compute_expenditure_amt', copy=False)
    income_source_ids = fields.One2many('plan.funding.source', 'plan_id', string="Source of Income", copy=False)
#     expenditure_ids = fields.One2many('plan.expenditure','plan_id', string="Expenditure", copy=False)
    note = fields.Text(string="Note", copy=False)
    section1_ids = fields.One2many('plan.section1', 'plan_id', string="Section 1", copy=False)
    section2_ids = fields.One2many('plan.section2', 'plan_id', string="Section 2", copy=False)
    section3_ids = fields.One2many('plan.section3', 'plan_id', string="Section 3", copy=False)
    section4_ids = fields.One2many('plan.section4', 'plan_id', string="Section 4", copy=False)
    show_section = fields.Char(compute='_compute_show_section')
    attachment_id = fields.Many2one('ir.attachment', string="MOM Document")
    state = fields.Selection([('draft', 'Draft'),('open', 'Open'),('confirm', 'Confirmed'),('approve', 'Approved'),('close', 'Closed')], string="Status", default='draft')
    user_id = fields.Many2one('res.users', string="User", default=lambda self: self.env.user)
    currency_id = fields.Many2one('res.currency', string='Currency', readonly=True, default=lambda self: self.env.user.company_id.currency_id.id)
    
    def action_draft(self):
        self.state = 'draft'
        
    def action_open(self):
        self.state = 'open'
        
    def action_confirm(self):
        self.state = 'confirm'
        
    def action_approve(self):
        self.state = 'approve'
        
    def action_close(self):
        self.state = 'close'
    
    @api.depends('section1','section2','section3','section4')
    def _compute_show_section(self):
        self.show_section = False
        for rec in self:
            if rec.section1 == 'yes':
                rec.show_section = 'sec1'
            elif rec.section2 == 'yes':
                rec.show_section = 'sec2'
            elif rec.section3 == 'yes':
                rec.show_section = 'sec3'
            elif rec.section4 == 'yes':
                rec.show_section = 'sec4'
    
    def open_all_activities(self):
        action = self.env.ref('planner.action_plan_activity').read()[0]
        action.update({
            'domain': [('plan_id','=',self.id)],
            'context': "{'act_view':1}",
        })
        return action
    
    def open_all_expenditure(self):
        action = self.env.ref('planner.action_plan_expenditure').read()[0]
        action.update({
            'domain': [('plan_id','=',self.id)],
            'context': "{'act_view':1,'default_plan_id':%d}"%(self.id),
        })
        return action
    
    def open_main_section(self):
        sec1_config_id = self.env['plan.section.config'].search([('section_type','=','section 1'),('user_ids','=',self.env.user.id)])
        sec2_config_id = self.env['plan.section.config'].search([('section_type','=','section 2'),('user_ids','=',self.env.user.id)])
        sec3_config_id = self.env['plan.section.config'].search([('section_type','=','section 3'),('user_ids','=',self.env.user.id)])
        sec4_config_id = self.env['plan.section.config'].search([('section_type','=','section 4'),('user_ids','=',self.env.user.id)])
        if self.show_section == 'sec1':
            action = self.env.ref('planner.action_plan_section1').read()[0]
            action.update({
                'display_name': sec1_config_id.name,
                'domain': [('plan_id','=',self.id)],
                'context': "{'default_plan_id':%d}"%(self.id),
            })
            return action
        elif self.show_section == 'sec2':
            action = self.env.ref('planner.action_plan_section2').read()[0]
            action.update({
                'display_name': sec2_config_id.name,
                'domain': [('plan_id','=',self.id)],
                'context': "{'default_plan_id':%d}"%(self.id),
            })
            return action
        elif self.show_section == 'sec3':
            action = self.env.ref('planner.action_plan_section3').read()[0]
            action.update({
                'display_name': sec3_config_id.name,
                'domain': [('plan_id','=',self.id)],
                'context': "{'default_plan_id':%d}"%(self.id),
            })
            return action
        elif self.show_section == 'sec4':
            action = self.env.ref('planner.action_plan_section4').read()[0]
            action.update({
                'display_name': sec4_config_id.name,
                'domain': [('plan_id','=',self.id)],
                'context': "{'default_plan_id':%d}"%(self.id),
            })
            return action
        else:
            raise Warning(_('Sorry! Choose the section correctly to open.'))

    def get_activities_data(self, document, field_cmp, section):
        sub2 = document.add_paragraph()
        sub2.add_run('ACTIVITIES').bold = True
        sub2.paragraph_format.left_indent = Pt(20)
        activity_ids = self.env['plan.activity'].search(
            [('main_activity_type_id', 'in', self.activity_type_ids.ids), (field_cmp, '=', section.id)])
        if activity_ids:
            table = document.add_table(rows=len(activity_ids) + 1, cols=5, style="Light Shading Accent 1")
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Activity Name'
            hdr_cells[0].width = Inches(3)
            hdr_cells[0].autofit = True
            hdr_cells[1].text = 'Activity Type'
            hdr_cells[0].width = Inches(2)
            hdr_cells[2].text = 'Responsible'
            hdr_cells[3].text = 'Status'
            hdr_cells[4].text = 'Analysis Status'
            row = 1
            for activity_id in activity_ids:
                cells = table.rows[row].cells
                cells[0].text = activity_id.name
                cells[1].text = activity_id.main_activity_type_id.name
                cells[2].text = activity_id.responsible_id.full_name
                if activity_id.status:
                    cells[3].text = dict(activity_id._fields['status'].selection).get(activity_id.status)
                if activity_id.analysis_status:
                    cells[4].text = dict(activity_id._fields['analysis_status'].selection).get(
                        activity_id.analysis_status)
                row = row + 1
    
    def generate_mom_report(self):
        document = Document()
        document.add_heading(self.name, 0)
        
        date_from = document.add_paragraph()
        date_from.add_run('Start Date: ').bold = True
        date_from.add_run(str(self.start_date.strftime("%d-%b-%Y")))
        date_to = document.add_paragraph()
        date_to.add_run('End Date: ').bold = True
        date_to.add_run(str(self.end_date.strftime("%d-%b-%Y")))
        category = document.add_paragraph()
        category.add_run('Category: ').bold = True
        category.add_run(self.category_id.name)
        responsible = document.add_paragraph()
        responsible.add_run('Primarily Responsible: ').bold = True
        responsible.add_run(self.responsible_id.display_name)
        
        document.add_page_break()
        
        current_user = self.env.user
        section1_config_id = self.env['plan.section.config'].search([('section_type','=','section 1'),('user_ids','=',current_user.id)])
        section2_config_id = self.env['plan.section.config'].search([('section_type','=','section 2'),('user_ids','=',current_user.id)])
        section3_config_id = self.env['plan.section.config'].search([('section_type','=','section 3'),('user_ids','=',current_user.id)])
        section4_config_id = self.env['plan.section.config'].search([('section_type','=','section 4'),('user_ids','=',current_user.id)])
        
        if self.section1 == 'yes':
            for section in self.section1_ids:
                sec1_head = document.add_heading(section.name + " (" + (section1_config_id.name if section1_config_id else 'DIMENSION') + ")",  level=1)
                sub = document.add_paragraph()
                sub.add_run('\n')
                sub.paragraph_format.left_indent = Pt(20)
                sec2 = document.add_paragraph()
                if self.section2 == 'yes':
                    sec2.add_run(section2_config_id.name if section2_config_id else 'NEED').bold = True
                    sec2.paragraph_format.left_indent = Pt(20)
                    numId = document.get_new_list("1")
                    for ct in section.section2_ids.mapped('name'):
                        p = document.add_paragraph(ct,style='List Number')
                        p.paragraph_format.left_indent = Pt(45)
                        p.num_id = numId
                sec3 = document.add_paragraph()
                if self.section3 == 'yes':
                    sec3.add_run(section3_config_id.name if section3_config_id else 'OBJECTIVE').bold = True
                    sec3.paragraph_format.left_indent = Pt(20)
                    numId = document.get_new_list("1")
                    for ct in section.section3_ids.mapped('name'):
                        p = document.add_paragraph(ct,style='List Number')
                        p.paragraph_format.left_indent = Pt(45)
                        p.num_id = numId
                sub1 = document.add_paragraph()
                if self.section4 == 'yes':
                    sub1.add_run(section4_config_id.name if section4_config_id else 'STRATEGY').bold = True
                    sub1.paragraph_format.left_indent = Pt(20)
                    numId = document.get_new_list("1")
                    for ct in section.section4_ids.mapped('name'):
                        p = document.add_paragraph(ct,style='List Number')
                        p.paragraph_format.left_indent = Pt(45)
                        p.num_id = numId
                self.get_activities_data(document,'section1_id', section)
         
        if self.section2 == 'yes' and self.section1 == 'no':
            for section in self.section2_ids:
                sec1_head = document.add_heading(section.name + " (" + (section2_config_id.name if section2_config_id else 'NEED') + ")",  level=1)
                sub = document.add_paragraph()
                sub.add_run('\n')
                sub.paragraph_format.left_indent = Pt(20)
                if self.section3 == 'yes':
                    sub.add_run(section3_config_id.name if section3_config_id else 'OBJECTIVE').bold = True
                    numId = document.get_new_list("1")
                    for ct in section.section3_ids.mapped('name'):
                        p = document.add_paragraph(ct,style='List Number')
                        p.paragraph_format.left_indent = Pt(45)
                        p.num_id = numId
                sub1 = document.add_paragraph()
                if self.section4 == 'yes':
                    sub1.add_run(section4_config_id.name if section4_config_id else 'STRATEGY' ).bold = True
                    sub1.paragraph_format.left_indent = Pt(20)
                    numId = document.get_new_list("1")
                    for ct in section.section4_ids.mapped('name'):
                        p = document.add_paragraph(ct,style='List Number')
                        p.paragraph_format.left_indent = Pt(45)
                        p.num_id = numId
                self.get_activities_data(document,'section2_id', section)
                        
        if self.section3 == 'yes' and self.section2 == 'no' and self.section1 == 'no':
            for section in self.section3_ids:
                sec1_head = document.add_heading(section.name + " (" + (section3_config_id.name if section3_config_id else 'OBJECTIVE') + ")",  level=1)
                sub = document.add_paragraph()
                sub.add_run('\n')
                sub.paragraph_format.left_indent = Pt(20)
                sub1 = document.add_paragraph()
                if self.section4 == 'yes':
                    sub1.add_run(section4_config_id.name if section4_config_id else 'STRATEGY').bold = True
                    sub1.paragraph_format.left_indent = Pt(20)
                    numId = document.get_new_list("1")
                    for ct in section.section4_ids.mapped('name'):
                        p = document.add_paragraph(ct,style='List Number')
                        p.paragraph_format.left_indent = Pt(45)
                        p.num_id = numId
                
                self.get_activities_data(document,'section3_id', section)
                
        if self.section4 == 'yes' and self.section3 == 'no' and self.section2 == 'no' and self.section1 == 'no':
            for section in self.section4_ids:
                sec1_head = document.add_heading(section.name + " (" + (section4_config_id.name if section4_config_id else 'STRATEGY') + ")",  level=1)
                sub = document.add_paragraph()
                sub.add_run('\n')
                sub.paragraph_format.left_indent = Pt(20)
                
                self.get_activities_data(document,'section4_id', section)
                    
        docx_stream = BytesIO()
        document.save(docx_stream)
        docx_bytes = docx_stream.getvalue()
        file = base64.encodestring(docx_bytes)
        
        file_name = self.name + '.docx'
        attachment = self.env['ir.attachment'].create({
                    'datas': file,
                    'type': 'binary',
                    'res_model': 'project.plan',
                    'db_datas': file_name,
                    'name': file_name,
                    })
        self.sudo().attachment_id = attachment.id

        base_url = self.env['ir.config_parameter'].sudo().get_param('web.base.url')
        return {
            'type': 'ir.actions.act_url',
            'target': 'self',
            'url': "{}/web/content?id={}&field=datas&model=ir.attachment&filename_field=name&download=true".format(base_url,self.attachment_id.id),
        }

class PlanSection1(models.Model):
    _name = "plan.section1"
    _inherit = 'common.plan.section'
    _description = 'Plan Section 1'
    
    @api.model
    def fields_view_get(self, view_id=None, view_type='form', toolbar=False, submenu=False):
        res = super(PlanSection1, self).fields_view_get(view_id=view_id, view_type=view_type, toolbar=toolbar, submenu=submenu)
        doc = etree.XML(res['arch'])
        current_user = self.env.user
        section2_config_id = self.env['plan.section.config'].search([('section_type','=','section 2'),('user_ids','=',current_user.id)])
        if section2_config_id:
            for node in doc.xpath("//notebook/page[3]"):
                if node:
                    node.set('string', section2_config_id.name)
            for node in doc.xpath("//button[@id='sec2']"):
                node.set('string', section2_config_id.name)
                        
        section3_config_id = self.env['plan.section.config'].search([('section_type','=','section 3'),('user_ids','=',current_user.id)])
        if section3_config_id:
            for node in doc.xpath("//notebook/page[4]"):
                if node:
                    node.set('string', section3_config_id.name)
            for node in doc.xpath("//button[@id='sec3']"):
                node.set('string', section3_config_id.name)
                
        section4_config_id = self.env['plan.section.config'].search([('section_type','=','section 4'),('user_ids','=',current_user.id)])
        if section4_config_id:
            for node in doc.xpath("//notebook/page[5]"):
                if node:
                    node.set('string', section4_config_id.name)
            for node in doc.xpath("//button[@id='sec4']"):
                node.set('string', section4_config_id.name)
        res['arch'] = etree.tostring(doc, encoding='unicode')
        return res
    
    name = fields.Char(string="Short Name", required=True)
    description = fields.Html(string="Description")
    activity_ids = fields.One2many('plan.activity', 'section1_id', string="Activity")
    section2_ids = fields.One2many('plan.section2', 'section1_id', string="Needs")
    section3_ids = fields.One2many('plan.section3', 'section1_id', string="Objective")
    section4_ids = fields.One2many('plan.section4', 'section1_id', string="Strategy")
    
    def open_activity(self):
        action = self.env.ref('planner.action_plan_activity').read()[0]
        action.update({
            'domain': [('section1_id','=',self.id)],
            'context': "{'default_plan_id':%d,'default_section1_id':%d}"%(self.plan_id.id,self.id),
        })
        return action
    
    def open_from_sec1(self):
        sec2_config_id = self.env['plan.section.config'].search([('section_type','=','section 2'),('user_ids','=',self.env.user.id)])
        sec3_config_id = self.env['plan.section.config'].search([('section_type','=','section 3'),('user_ids','=',self.env.user.id)])
        sec4_config_id = self.env['plan.section.config'].search([('section_type','=','section 4'),('user_ids','=',self.env.user.id)])
        if self._context.get('open_sec',False) == 'sec2':
            action = self.env.ref('planner.action_plan_section2').read()[0]
            action.update({
                'display_name': sec2_config_id.name,
                'domain': [('section1_id','=',self.id)],
                'context': "{'default_plan_id':%d,'default_section1_id':%d}"%(self.plan_id.id,self.id),
            })
            return action
        elif self._context.get('open_sec',False) == 'sec3':
            action = self.env.ref('planner.action_plan_section3').read()[0]
            action.update({
                'display_name': sec3_config_id.name,
                'domain': [('section1_id','=',self.id)],
                'context': "{'default_plan_id':%d,'default_section1_id':%d}"%(self.plan_id.id,self.id),
            })
            return action
        elif self._context.get('open_sec',False) == 'sec4':
            action = self.env.ref('planner.action_plan_section4').read()[0]
            action.update({
                'display_name': sec4_config_id.name,
                'domain': [('section1_id','=',self.id)],
                'context': "{'default_plan_id':%d,'default_section1_id':%d}"%(self.plan_id.id,self.id),
            })
            return action
        else:
            raise Warning(_('Sorry! Choose the section correctly to open.'))
        
class PlanSection2(models.Model):
    _name = "plan.section2"
    _inherit = 'common.plan.section'
    _description = 'Plan Section 2'
    
    @api.model
    def fields_view_get(self, view_id=None, view_type='form', toolbar=False, submenu=False):
        res = super(PlanSection2, self).fields_view_get(view_id=view_id, view_type=view_type, toolbar=toolbar, submenu=submenu)
        doc = etree.XML(res['arch'])
        current_user = self.env.user
        section3_config_id = self.env['plan.section.config'].search([('section_type','=','section 3'),('user_ids','=',current_user.id)])
        if section3_config_id:
            for node in doc.xpath("//notebook/page[3]"):
                if node:
                    node.set('string', section3_config_id.name)
            for node in doc.xpath("//button[@id='sec3']"):
                node.set('string', section3_config_id.name)
                
        section4_config_id = self.env['plan.section.config'].search([('section_type','=','section 4'),('user_ids','=',current_user.id)])
        if section4_config_id:
            for node in doc.xpath("//notebook/page[4]"):
                if node:
                    node.set('string', section4_config_id.name)
            for node in doc.xpath("//button[@id='sec4']"):
                node.set('string', section4_config_id.name)
        res['arch'] = etree.tostring(doc, encoding='unicode')
        return res
    
    name = fields.Char(string="Short Name", required=True)
    description = fields.Html(string="Description")
    activity_ids = fields.One2many('plan.activity', 'section2_id', string="Activity")
    section3_ids = fields.One2many('plan.section3', 'section2_id', string="Objective")
    section4_ids = fields.One2many('plan.section4', 'section2_id', string="Strategy")
#     plan_id = fields.Many2one('project.plan', string="Project Plan", ondelete='cascade')
    section1_id = fields.Many2one('plan.section1', string="Dimension", ondelete='cascade')
    
    def open_activity(self):
        action = self.env.ref('planner.action_plan_activity').read()[0]
        action.update({
            'domain': [('section2_id','=',self.id)],
            'context': "{'default_plan_id':%d,'default_section1_id':%d,'default_section2_id':%d}"%(self.plan_id.id,self.section1_id.id,self.id),
        })
        return action
    
    def open_from_sec2(self):
        sec3_config_id = self.env['plan.section.config'].search([('section_type','=','section 3'),('user_ids','=',self.env.user.id)])
        sec4_config_id = self.env['plan.section.config'].search([('section_type','=','section 4'),('user_ids','=',self.env.user.id)])
        if self._context.get('open_sec',False) == 'sec3':
            action = self.env.ref('planner.action_plan_section3').read()[0]
            action.update({
                'display_name': sec3_config_id.name,
                'domain': [('section2_id','=',self.id)],
                'context': "{'default_plan_id':%d,'default_section1_id':%d,'default_section2_id':%d}"%(self.plan_id.id,self.section1_id.id,self.id),
            })
            return action
        elif self._context.get('open_sec',False) == 'sec4':
            action = self.env.ref('planner.action_plan_section4').read()[0]
            action.update({
                'display_name': sec4_config_id.name,
                'domain': [('section2_id','=',self.id)],
                'context': "{'default_plan_id':%d,'default_section1_id':%d,'default_section2_id':%d}"%(self.plan_id.id,self.section1_id.id,self.id),
            })
            return action
        else:
            raise Warning(_('Sorry! Choose the section correctly to open.'))
    
class PlanSection3(models.Model):
    _name = "plan.section3"
    _inherit = 'common.plan.section'
    _description = 'Plan Section 3'
    
    @api.model
    def fields_view_get(self, view_id=None, view_type='form', toolbar=False, submenu=False):
        res = super(PlanSection3, self).fields_view_get(view_id=view_id, view_type=view_type, toolbar=toolbar, submenu=submenu)
        doc = etree.XML(res['arch'])
        current_user = self.env.user
        section4_config_id = self.env['plan.section.config'].search([('section_type','=','section 4'),('user_ids','=',current_user.id)])
        if section4_config_id:
            for node in doc.xpath("//notebook/page[3]"):
                if node:
                    node.set('string', section4_config_id.name)
            for node in doc.xpath("//button[@id='sec4']"):
                node.set('string', section4_config_id.name)
        res['arch'] = etree.tostring(doc, encoding='unicode')
        return res
    
    name = fields.Char(string="Short Name", required=True)
    description = fields.Html(string="Description")
    activity_ids = fields.One2many('plan.activity', 'section3_id', string="Activity")
    section4_ids = fields.One2many('plan.section4', 'section3_id', string="Strategy")
#     plan_id = fields.Many2one('project.plan', string="Project Plan", ondelete='cascade')
    section1_id = fields.Many2one('plan.section1', string="Dimension", ondelete='cascade')
    section2_id = fields.Many2one('plan.section2', string="Need", ondelete='cascade')
    
    def open_activity(self):
        action = self.env.ref('planner.action_plan_activity').read()[0]
        action.update({
            'domain': [('section3_id','=',self.id)],
            'context': "{'default_plan_id':%d,'default_section1_id':%d,'default_section2_id':%d,'default_section3_id':%d}"%(self.plan_id.id,self.section1_id.id,self.section2_id.id,self.id),
        })
        return action
    
    def open_from_sec3(self):
        sec4_config_id = self.env['plan.section.config'].search([('section_type','=','section 4'),('user_ids','=',self.env.user.id)])
        if self._context.get('open_sec',False) == 'sec4':
            action = self.env.ref('planner.action_plan_section4').read()[0]
            action.update({
                'display_name': sec4_config_id.name,
                'domain': [('section3_id','=',self.id)],
                'context': "{'default_plan_id':%d,'default_section1_id':%d,'default_section2_id':%d,'default_section3_id':%d}"%(self.plan_id.id,self.section1_id.id,self.section2_id.id,self.id),
            })
            return action
        else:
            raise Warning(_('Sorry! Choose the section correctly to open.'))
    
class PlanSection4(models.Model):
    _name = "plan.section4"
    _inherit = 'common.plan.section'
    _description = 'Plan Section 4'
    
    name = fields.Char(string="Short Name", required=True)
    description = fields.Html(string="Description")
    activity_ids = fields.One2many('plan.activity', 'section4_id', string="Activity")
#     plan_id = fields.Many2one('project.plan', string="Project Plan", ondelete='cascade')
    section1_id = fields.Many2one('plan.section1', string="Dimension", ondelete='cascade')
    section2_id = fields.Many2one('plan.section2', string="Need", ondelete='cascade')
    section3_id = fields.Many2one('plan.section3', string="Objective", ondelete='cascade') 
    
    def open_activity(self):
        action = self.env.ref('planner.action_plan_activity').read()[0]
        action.update({
            'domain': [('section4_id','=',self.id)],
            'context': "{'default_plan_id':%d,'default_section1_id':%d,'default_section2_id':%d,'default_section3_id':%d,'default_section4_id':%d}"%(self.plan_id.id,self.section1_id.id,self.section2_id.id,self.section3_id.id,self.id),
        })
        return action
    
# class PlanActivityType(models.Model):
#     _name = "plan.activity.type"
#     _description = "Plan Activity Type"
#     _rec_name = "main_activity_type_id"
#     
#     main_activity_type_id = fields.Many2one('plan.main.activity.type', string="Main Activity Type", required=True)
#     activity_ids = fields.One2many('plan.activity', 'activity_type_id', string="Activities")
#     plan_id = fields.Many2one('project.plan',string="Plan", ondelete='cascade')
#     section1_id = fields.Many2one('plan.section1', string="Dimension", ondelete='cascade')
#     section2_id = fields.Many2one('plan.section2', string="Need", ondelete='cascade')
#     section3_id = fields.Many2one('plan.section3', string="Objective", ondelete='cascade')
#     section4_id = fields.Many2one('plan.section4', string="Strategy", ondelete='cascade')
#     plan_act_type_ids = fields.Many2many('plan.main.activity.type', related='plan_id.activity_type_ids', string="Plan Activity Type")
#     
#     def open_activities(self):
#         action = self.env.ref('planner.action_plan_activity').read()[0]
#         action.update({
#             'domain': [('activity_type_id','=',self.id)],
#             'context': "{'default_plan_id':%d,'default_activity_type_id':%d}"%(self.plan_id.id,self.id),
#         })
#         return action
    
class PlanActivity(models.Model):
    _name = "plan.activity"
    _description = "Plan Activity"
    
    @api.model
    def fields_view_get(self, view_id=None, view_type='form', toolbar=False, submenu=False):
        res = super(PlanActivity, self).fields_view_get(view_id=view_id, view_type=view_type, toolbar=toolbar, submenu=submenu)
        doc = etree.XML(res['arch'])
        current_user = self.env.user
        if self._context.get('act_view'):
            doc.set('create','false')
        res['arch'] = etree.tostring(doc, encoding='unicode')
        return res
    
    def _compute_expenditure_amt(self):
        self.actual_expenditure = 0
        for rec in self:
            for expenditure in rec.expenditure_ids:
                rec.actual_expenditure +=  expenditure.amount
    
    name = fields.Char(string="Activity Name", required=True) 
    section1_id = fields.Many2one('plan.section1', string="Section 1", ondelete='cascade')
    section2_id = fields.Many2one('plan.section2', string="Section 2", ondelete='cascade')
    section3_id = fields.Many2one('plan.section3', string="Section 3", ondelete='cascade')
    section4_id = fields.Many2one('plan.section4', string="Section 4", ondelete='cascade')
    main_activity_type_id = fields.Many2one('plan.main.activity.type', string="Activity Type",required=True)
    start_date = fields.Date(string="Start Date")
    end_date = fields.Date(string="End Date")
    responsible_id = fields.Many2one('res.partner', string="Responsible")
    details = fields.Html(string="Activity Details")
    expected_outcome = fields.Char(string="Expected Outcome") 
    status = fields.Selection([('yet to start','Yet to Start'),('inprogress','In-progress'),('completed','Completed')], string="Status",default='yet to start')
    evaluation_criteria = fields.Text(string="Evaluation Criteria")
    analysis_status = fields.Selection([('open','Open'),('inprogress','Inprogress'),('done','Done')], string="Analysis Status")
    expenditure_ids = fields.One2many('plan.expenditure','activity_id', string="Expenditure")
    plan_id = fields.Many2one('project.plan', string="Plan", ondelete='cascade')
    actual_expenditure = fields.Float(string="Actual Expenditure", compute='_compute_expenditure_amt', copy=False)
    currency_id = fields.Many2one('res.currency', string='Currency', readonly=True, default=lambda self: self.env.user.company_id.currency_id.id)


    @api.constrains('start_date','end_date')
    def validate_date(self):
        if self.start_date:
            if (self.start_date < self.plan_id.start_date or self.start_date > self.plan_id.end_date):
                raise UserError(_("Activity Start date must be Within the Plan Start date and End date."))
        if self.end_date:
            if (self.end_date > self.plan_id.end_date or self.end_date < self.plan_id.start_date):
                raise UserError(_("Activity End date must be Within the Plan Start date and End date."))
        if self.start_date and self.end_date:
            if self.start_date > self.end_date:
                raise UserError(_("Start date must be less than the End date."))
            
        
class PlanExpenditure(models.Model):
    _name = "plan.expenditure"
    _description = "Plan Expenditure" 
    _rec_name = "activity_id"

    # @api.model
    # def fields_view_get(self, view_id=None, view_type='form', toolbar=False, submenu=False):
    #     res = super(PlanExpenditure, self).fields_view_get(view_id=view_id, view_type=view_type, toolbar=toolbar, submenu=submenu)
    #     doc = etree.XML(res['arch'])
    #     print(self._context)
    #     if self._context.get('act_view'):
    #         nodes = doc.xpath("//field[@name='activity_id']")
    #         for node in nodes:
    #             if self._context.get('plan_id'):
    #                 node.set('domain', "[('plan_id', '=', %d)]"%(self._context.get('plan_id')))
    #     res['arch'] = etree.tostring(doc, encoding='unicode')
    #     return res
    
    activity_id = fields.Many2one('plan.activity', string="Activity Reference", ondelete='cascade', required=True)
    particular = fields.Char(string="Material/Service")
    quantity = fields.Float(string="Quantity")
    amount = fields.Float(string="Amount")
    expenditure_item_ids = fields.One2many('plan.expenditure.item','expenditure_id', string="Expenditure")
    plan_id = fields.Many2one('project.plan', string="Plan", ondelete='cascade')
    currency_id = fields.Many2one('res.currency', string='Currency', readonly=True, default=lambda self: self.env.user.company_id.currency_id.id)

    @api.constrains('activity_id')
    def update_plan_id(self):
        for rec in self:
            if rec.activity_id:
                rec.plan_id = rec.activity_id.plan_id

    @api.model
    def create(self, vals):
        if self._context.get('import_file'):
            if vals.get('plan_id'):
                activity_ids = self.env['plan.activity'].search([('plan_id','=',vals.get('plan_id'))]).ids
                if not vals.get('activity_id') in activity_ids:
                    raise UserError(_("Sorry! The given activity is not belong to the given plan."))
            else:
                raise UserError(_("Sorry! Plan name is must"))
        return super(PlanExpenditure, self).create(vals)

class PlanExpenditureItem(models.Model):
    _name = "plan.expenditure.item"
    _description = "Plan Expenditure Item" 
    _rec_name = "item"
    
    expenditure_date = fields.Date(string="Expenditure Date", required=True)
    item = fields.Char(string="Expenditure Item", required=True)
    expenditure_id = fields.Many2one('plan.expenditure', string="Expenditure", ondelete='cascade')
    
    @api.constrains('expenditure_date')
    def validate_expenditure_date(self):
        for rec in self:
            if rec.expenditure_date:
                if rec.expenditure_date < rec.expenditure_id.plan_id.start_date:
                    raise UserError(_("Expenditure date must be Within the plan Start date."))
                if rec.expenditure_id.plan_id.end_date and rec.expenditure_date > rec.expenditure_id.plan_id.end_date:
                    raise UserError(_("Expenditure date must be Within the plan End date."))
    
class PlanFundingSource(models.Model):
    _name = "plan.funding.source"
    _description = "Plan Funding Source" 
    _rec_name = "source_name"
    
    source_name = fields.Char(string="Source Name", required=True)
    source_type_id = fields.Many2one('fund.source.type', string="Source Type")
    amount_identified = fields.Float(string="Amount Identified")
    plan_id = fields.Many2one('project.plan', string="Project Plan", ondelete='cascade')
    currency_id = fields.Many2one('res.currency', string='Currency', readonly=True, default=lambda self: self.env.user.company_id.currency_id.id)
    