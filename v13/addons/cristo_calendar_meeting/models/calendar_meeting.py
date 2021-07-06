# -*- coding: utf-8 -*-
from odoo import models, fields, api, tools, _
from odoo.exceptions import UserError
from odoo.exceptions import AccessError, UserError, ValidationError
from odoo.addons.cristo.tools import cris_tools
from odoo.addons.calendar.models import calendar
# from odoo.tools.misc import formatLang
from docx import Document
from docx.shared import Pt
import base64
from io import BytesIO
import html2text
from dateutil import tz
from datetime import datetime, timedelta
from odoo.addons.cristo.models.res_common import Mail_Excluded_Fields

class CalendarMeeting(models.Model):
    _name = 'calendar.event'
    _inherit = ['calendar.event','common.rel.fields','attachment.size','common.ecc.fields']
    _description = "Calendar Meeting"
    _custom_filter_exclude_fields = ['org_image_ids','attachment_id','attachment_ids','description','activity_ids','message_ids',
                                     'attendee_ids','rrule','rrule_type','count','interval','final_date','alarm_ids','end_type','res_id',
                                     'res_model_id','res_model','recurrency','recurrent_id','recurrent_id_date','interval','show_as',
                                     'event_tz','month_by','proceedings','agenda','decision','description_html','byday','week_list',
                                     'mo','tu','we','th','fr','sa','su','day',
                                     ] + Mail_Excluded_Fields
    
    category = fields.Selection([('calendar','Calendar'),('meeting','Meeting')], string="Category")
    proceedings = fields.Html(string="Proceedings")
    agenda = fields.Html(string="Agenda")
    decision = fields.Text(string="Decision")
    attachment_ids = fields.Many2many('ir.attachment', string="Attach Files")
    email_to = fields.Char(string="Email To")
    email_cc = fields.Char(string="Email CC")
    attendee_type = fields.Selection([('individual','Individual'),('group','Group')], string="Attendee Type", default="individual")
    meeting_group_ids = fields.Many2many('meeting.group', string="Meeting Group")
    org_image_ids = fields.One2many('org.image','calendar_id', string="Extra Org Media", copy=True)
    calendar_type_id = fields.Many2one('calendar.type', string="Calendar Type")
    description_html = fields.Html(string="Description")
    partner_ids = fields.Many2many('res.partner', 'calendar_event_res_partner_rel', string='Attendees')
    main_category_id = fields.Many2one('res.main.category', string="Main Category")
    main_category_code = fields.Char(related="main_category_id.code", store=True)
    attachment_id = fields.Many2one('ir.attachment',string="MOM Document")
    type = fields.Selection([('public','Public'),('private','Private')], string="Type", default="private")
    calendar_type = fields.Selection([('province_calendar','Province Calendar'),('provincial_calendar','Provincial Calendar'),('community_calendar','Community Calendar'),('institution_calendar','Institution Calendar'),('personal_calendar','Personal Calendar')], string="Calendar Type")
    # user_id = fields.Many2one('res.users',compute="_compute_user",store=True,readonly=False)
    send_email = fields.Boolean(compute='_check_send_email')

    @api.depends('send_email')
    def _check_send_email(self):
        for rec in self:
            if self.env.user.id == rec.create_uid.id:
                rec.send_email = True
            else:
                rec.send_email = False
           
    @api.constrains('type')
    def _check_calendar_type(self):
        if self.category == 'calendar':
            if not self.type:
                raise ValidationError(_("Sorry! You must choose Type."))
    
    @api.constrains('attachment_ids')
    def _check_attachment_size(self):
        self.env['ir.attachment']._check_size(self.attachment_ids)
        
    @api.constrains('calendar_type')
    def _check_calendar_type(self):
        if self.user_has_groups('cristo.group_role_cristo_religious_province'):
            if self.calendar_type !='province_calendar' and self.calendar_type != 'provincial_calendar':
                raise ValidationError(_("Sorry! You must choose proper calendar type."))
            
        elif self.user_has_groups('cristo.group_role_cristo_religious_house'):
            if self.calendar_type != 'community_calendar':
                raise ValidationError(_("Sorry! You must choose proper calendar type."))
            
        elif self.user_has_groups('cristo.group_role_cristo_apostolic_institution'):
            if self.calendar_type != 'institution_calendar':
                raise ValidationError(_("Sorry! You must choose proper calendar type."))
            
        elif self.user_has_groups('cristo.group_role_cristo_individual'):
            if self.calendar_type != 'personal_calendar':
                raise ValidationError(_("Sorry! You must choose proper calendar type."))
    @api.model
    def default_get(self, fields):
        data = super(CalendarMeeting, self).default_get(fields)
        if self.user_has_groups('cristo.group_role_cristo_religious_province'):
            data['calendar_type'] = 'province_calendar'
        elif self.user_has_groups('cristo.group_role_cristo_religious_house'):
            data['calendar_type'] = 'community_calendar'
        elif self.user_has_groups('cristo.group_role_cristo_apostolic_institution'):
            data['calendar_type'] = 'institution_calendar'
        elif self.user_has_groups('cristo.group_role_cristo_individual'):
            data['calendar_type'] = 'personal_calendar'
        return data
        
    def convert_utc(self, utc_datetime):
        from_zone = tz.gettz('UTC')
        user_id = self.env['res.users'].browse(self.env.uid)
        to_zone = tz.gettz( user_id.tz or 'Asia/Calcutta')
        utc = datetime.strptime(utc_datetime, '%Y-%m-%d %H:%M:%S')
        utc = utc.replace(tzinfo=from_zone)
        central = utc.astimezone(to_zone)
        return central

    def generate_mom_report(self):
        document = Document()
#         style = document.styles['Normal']
#         font = style.font
#         font.name = 'DejaVu Serif'
#         font.size = Pt(11)
        
        if self.start_datetime:
            start_datetime = self.convert_utc(str(self.start_datetime))
            date_time = document.add_paragraph()
            date_time.add_run('Date & Time : ').bold = True
            date_time.add_run(str(start_datetime.strftime('%d/%m/%Y') + '\t'))
            date_time.add_run('@' + str(start_datetime.strftime('%H:%M %p')))
        else:
            date_time = document.add_paragraph()
            date_time.add_run('Date : ').bold = True
            date_time.add_run(str(self.start_date))

        present = []
        absent = []
        present_participant = document.add_paragraph()
        absent_participant = document.add_paragraph()
        present_participant.add_run('Present Participants : ').bold = True
        absent_participant.add_run('Absent Participants : ').bold = True
        if self.attendee_ids:
            for attendee in self.attendee_ids:
                if attendee.state == 'accepted':
                    present.append(attendee.partner_id.name)
                elif attendee.state == 'declined':
                    absent.append(attendee.partner_id.name)

        present = ", ".join(present) if present else "None"
        absent = ", ".join(absent) if absent else "None"
        present_participant.add_run(present)
        absent_participant.add_run(absent)
        
        
        proceedings = html2text.html2text(self.proceedings) if self.proceedings else ""
        delibration = document.add_paragraph()
        delibration.add_run('PROCEEDINGS').bold = True
        delibration.add_run('\n')
        delibration.add_run(proceedings)
           
        decision = document.add_paragraph()
        decision.add_run('DECISIONS').bold = True
        decision.add_run('\n')
        decision.add_run(self.decision)
        
        description = html2text.html2text(self.description_html) if self.description_html else ""
        action = document.add_paragraph()
        action.add_run('ACTIONS ITEMS').bold = True
        action.add_run('\n')
        action.add_run(description)
        action.add_run('==============================================')
#         document.add_page_break()
        
        docx_stream = BytesIO()
        document.save(docx_stream)
        docx_bytes = docx_stream.getvalue()
        file = base64.encodestring(docx_bytes)
        
        if self.attachment_id:
            self.attachment_id.write({'datas' : file,})
        else:
            if self.start_datetime:
                start_datetime = fields.Datetime.from_string(self.start_datetime)
                file_name = "MOM-CRISTO-" + start_datetime.strftime('%Y') + '-' + start_datetime.strftime('%B') + '.docx'
            else:
                file_name = "MOM-CRISTO-" + str(self.start_date) + '.docx'
            real_id = calendar.calendar_id2real_id(self.id)
            attachment = self.env['ir.attachment'].create({
                    'datas': file,
                    'type': 'binary',
                    'res_model': 'calendar.event',
                    'res_id': real_id,
                    'db_datas': file_name,
#                     'datas_fname':file_name,
                    'name':file_name,
                    })
            self.attachment_id = attachment.id

        base_url = self.env['ir.config_parameter'].sudo().get_param('web.base.url')
        return {
            'type': 'ir.actions.act_url',
            'target': 'self',
            'url': "{}/web/content?id={}&field=datas&model=ir.attachment&filename_field=name&download=true".format(base_url,self.attachment_id.id),
        }
            
#     This function is to open the appropriate form view from calendar edit button. This called from JS.
    def get_formview_id(self, access_uid=None):
        """ Override this method in order to redirect many2one towards the right model depending on access_uid """
        if self._context.get('default_category',False) == 'meeting':
            return self.env.ref('cristo_calendar_meeting.view_meeting_form').id
        elif (self._context.get('default_category',False) == 'calendar'):
            return self.env.ref('cristo_calendar_meeting.view_calendar_form').id
        else:
            return self.env.ref('calendar.view_calendar_event_form').id
    
    @api.onchange('start_datetime', 'duration')
    def _onchange_duration(self):
        if self.start_datetime:
            start = self.start_datetime
            self.start = self.start_datetime
            # Round the duration (in hours) to the minute to avoid weird situations where the event
            # stops at 4:19:59, later displayed as 4:19.
            self.stop = start + timedelta(minutes=round((self.duration or 1.0) * 60))
            if self.allday:
                self.stop -= timedelta(seconds=1)
    
    @api.constrains('email_to')
    def check_email_to(self):
        if self.email_to:
            cris_tools.email_validation(self.email_to)
            
    @api.constrains("email_cc")        
    def check_email_cc(self):
        if self.email_cc:
            cris_tools.email_validation(self.email_cc)
    
    @api.constrains('org_image_ids')
    def validate_images(self):
        if len(self.org_image_ids) > 5:
            raise ValidationError(_("Sorry! You can only add 5 media."))
    
    @api.onchange('attendee_type')
    def _onchange_attendee_type(self):
        if self.attendee_type == 'individual':
            self.meeting_group_ids = False
            self.partner_ids = self.env.user.partner_id
        if self.attendee_type == 'group':
            self.partner_ids = False
            
    def action_detach_recurring_event(self):
        meeting = self.detach_recurring_event()
        if meeting.category == 'meeting':
            view_id = self.env.ref('cristo_calendar_meeting.view_meeting_form').id
        elif meeting.category == 'calendar':
            view_id = self.env.ref('cristo_calendar_meeting.view_calendar_form').id
        return {
            'type': 'ir.actions.act_window',
            'res_model': 'calendar.event',
            'view_mode': 'form',
            'view_id': view_id,
            'res_id': meeting.id,
            'target': 'current',
            'flags': {'form': {'action_buttons': True, 'options': {'mode': 'edit'}}}
        }

    def action_meeting_wizard(self):
        action = self.env.ref('cristo_calendar_meeting.action_wizard_meeting').read()[0]
        if self.rrule:
            event_id = self.detach_recurring_event().id
        else:
            event_id = self.id
        action.update({
            'context': {'default_calendar_event_id':event_id,'default_body':str(self.proceedings) + str(self.agenda),'default_subject':'Invitation to '+self.name,'default_attachment_ids':[(6, 0, self.attachment_ids.ids)]}
            })
        return action
    
#     The Below function is to overwrite Attendee based on Attendee Type.This will be called in create and write method.
    def create_attendees(self):
        current_user = self.env.user
        result = {}
        for meeting in self:
            alreay_meeting_partners = meeting.attendee_ids.mapped('partner_id')
            meeting_attendees = self.env['calendar.attendee']
            meeting_partners = self.env['res.partner']
            if meeting.attendee_type == 'individual':
                for partner in meeting.partner_ids.filtered(lambda partner: partner not in alreay_meeting_partners):
                    values = {
                        'partner_id': partner.id,
                        'email': partner.email,
                        'event_id': meeting.id,
                    }
                    # current user don't have to accept his own meeting
                    if partner == self.env.user.partner_id:
                        values['state'] = 'accepted'
    
                    attendee = self.env['calendar.attendee'].create(values)
    
                    meeting_attendees |= attendee
                    meeting_partners |= partner
                    
            elif meeting.attendee_type == 'group':
                if meeting.meeting_group_ids:
                    partner_ids = meeting.meeting_group_ids.partner_ids.filtered(lambda partner: partner not in alreay_meeting_partners)
                    if current_user.partner_id not in partner_ids and current_user.partner_id not in alreay_meeting_partners:
                        partner_ids =  current_user.partner_id + partner_ids
                    for partner in partner_ids:
                        values = {
                            'partner_id': partner.id,
                            'email': partner.email,
                            'event_id': meeting.id,
                        }
                        # current user don't have to accept his own meeting
                        if partner == self.env.user.partner_id:
                            values['state'] = 'accepted'
                        attendee = self.env['calendar.attendee'].create(values)
                        
        
                        meeting_attendees |= attendee
                        meeting_partners |= partner
                           
            if meeting_attendees and not self._context.get('detaching'):
                to_notify = meeting_attendees.filtered(lambda a: a.email != current_user.email)
                to_notify._send_mail_to_attendees('calendar.calendar_template_meeting_invitation')

            if meeting_attendees:
                meeting.write({'attendee_ids': [(4, meeting_attendee.id) for meeting_attendee in meeting_attendees]})

            if meeting_partners:
                meeting.message_subscribe(partner_ids=meeting_partners.ids)
            # We remove old attendees who are not in partner_ids now.
            if meeting.attendee_type == 'individual':
                all_partners = meeting.partner_ids
            elif meeting.attendee_type == 'group':
                all_partners = current_user.partner_id + meeting.meeting_group_ids.partner_ids
                
            all_partner_attendees = meeting.attendee_ids.mapped('partner_id')
            old_attendees = meeting.attendee_ids
            partners_to_remove = all_partner_attendees + meeting_partners - all_partners

            attendees_to_remove = self.env["calendar.attendee"]
            if partners_to_remove:
                attendees_to_remove = self.env["calendar.attendee"].search([('partner_id', 'in', partners_to_remove.ids), ('event_id', '=', meeting.id)])
                attendees_to_remove.unlink()

            result[meeting.id] = {
                'new_attendees': meeting_attendees,
                'old_attendees': old_attendees,
                'removed_attendees': attendees_to_remove,
                'removed_partners': partners_to_remove
            }
        return result
    
    def write(self, values):
        for meeting in self:
            real_ids = []
            new_ids = []
            if not calendar.is_calendar_id(meeting.id):
                real_ids = [int(meeting.id)]
            else:
                real_event_id = calendar.calendar_id2real_id(meeting.id)
                
                blacklisted = any(key in values for key in ('start', 'stop', 'active'))
                if not values.get('recurrency', True) or not blacklisted:
                    real_ids = [real_event_id]
                else:
                    data = meeting.read(['start', 'stop', 'rrule', 'duration'])[0]
                    if data.get('rrule'):
                        new_ids = meeting.with_context(dont_notify=True).detach_recurring_event(values).ids
            new_meetings = self.browse(new_ids)
            real_meetings = self.browse(real_ids)
            all_meetings = real_meetings + new_meetings
            super(CalendarMeeting, real_meetings).write(values)  
            
            if values.get('meeting_group_ids', False):
                attendees_create = all_meetings.with_context(dont_notify=True).create_attendees()
        return True
                       
class MeetingGroups(models.Model):
    _name = 'meeting.group'
    _description = "Meeting Groups"
    
    name = fields.Char(string="Name", required=True)
    partner_ids = fields.Many2many('res.partner', string="Attendees")
    user_id = fields.Many2one('res.users', default=lambda self: self.env.user)
    
    @api.constrains('partner_ids')
    def check_member(self):
        if not self.partner_ids:
            raise UserError(_("Select the Attendees."))
        
class CalendarType(models.Model):
    _name = "calendar.type"
    _description = "Calendar Type"
    
    name = fields.Char("Name", required=True)
    
